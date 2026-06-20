"""PDF- und DOCX-Export für gespeicherte Tagesberichte (Freiraum Baustellen-Doku)."""

from __future__ import annotations

import re
import unicodedata
from io import BytesIO
from pathlib import Path
from typing import Any, Callable
from urllib.parse import urlparse

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

GREY_META_HEX = HexColor("#52525b")
LINE_HEX = HexColor("#d4d4d8")
SOFT_BG_HEX = HexColor("#f4f4f5")
TEXT_DARK_HEX = HexColor("#111827")
SECTION_HEX = HexColor("#1f2937")
LOGO_DIR = Path(__file__).resolve().parent / "uploads" / "logos"
SIGNATURES_DIR = Path(__file__).resolve().parent / "uploads" / "signatures"
TENANTS_UPLOADS_DIR = Path(__file__).resolve().parent / "uploads" / "tenants"
SIGNATURE_ROLES = ("customer", "employee")

LogoPathResolver = Callable[[dict[str, Any], dict[str, Any]], Path | None]
SignaturePathResolver = Callable[[str], Path | None]


def sanitize_export_slug(s: str) -> str:
    t = (s or "").strip()
    if not t:
        return "Baustelle"
    t = re.sub(r"\s+", "_", t)
    # Wort-Zeichen inkl. Umlaute, Rest → _
    t = re.sub(r"[^\w\-]+", "_", t, flags=re.UNICODE)
    t = re.sub(r"_+", "_", t).strip("_")
    return t[:80] if t else "Baustelle"


def sanitize_export_slug_ascii(s: str) -> str:
    n = unicodedata.normalize("NFKD", (s or "").strip())
    n = n.encode("ascii", "ignore").decode("ascii")
    n = re.sub(r"\s+", "_", n)
    n = re.sub(r"[^a-zA-Z0-9_\-]", "_", n)
    n = re.sub(r"_+", "_", n).strip("_")
    return n[:80] if n else "Baustelle"


def build_export_base_name(report: dict[str, Any]) -> str:
    site = sanitize_export_slug(str(report.get("projectName") or "Baustelle"))
    d_raw = sanitize_export_slug(str(report.get("date") or "datum"))
    return f"tagesbericht_{site}_{d_raw}"


def build_attachment_names(report: dict[str, Any], ext: str) -> tuple[str, str]:
    ascii_nm = (
        "tagesbericht_"
        f"{sanitize_export_slug_ascii(str(report.get('projectName') or 'Baustelle')).lower()}_"
        f"{sanitize_export_slug_ascii(str(report.get('date') or 'datum')).lower()}.{ext}"
    )
    desc = f"{build_export_base_name(report)}.{ext}"
    return ascii_nm, desc


def _xml_para_text(s: str) -> str:
    t = s or ""
    return (
        t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
    )


def _structured(report: dict[str, Any]) -> dict[str, Any]:
    st = report.get("structured") or {}
    return st if isinstance(st, dict) else {}


def _list_or_keine(items: Any) -> list[str]:
    if isinstance(items, list) and items:
        return [str(x) for x in items]
    return ["Keine Angabe"]


def _format_date_de(date_str: str) -> str:
    """ISO YYYY-MM-DD → TT.MM.JJJJ für Export; sonst Originaltext."""
    s = (date_str or "").strip()
    if not s or s == "—":
        return s or "—"
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})", s)
    if m:
        return f"{m.group(3)}.{m.group(2)}.{m.group(1)}"
    return s


def _candidate_logo_paths(report: dict[str, Any], company_profile: dict[str, Any]) -> list[Path]:
    out: list[Path] = []

    logo_url = str(report.get("companyLogoUrl") or "").strip()
    if logo_url:
        parsed = urlparse(logo_url)
        path_only = parsed.path if parsed.path else logo_url
        marker = "/uploads/logos/"
        if marker in path_only:
            filename = Path(path_only.split(marker, 1)[1]).name
            if filename:
                out.append(LOGO_DIR / filename)
        tenant_marker = "/uploads/tenants/"
        if tenant_marker in path_only:
            rest = path_only.split(tenant_marker, 1)[1]
            parts = rest.strip("/").split("/")
            if len(parts) >= 3 and parts[1] == "logos":
                out.append(
                    Path(__file__).resolve().parent / "uploads" / "tenants" / parts[0] / "logos" / parts[2]
                )

    logo_fn = str(company_profile.get("logoFilename") or "").strip()
    if logo_fn:
        out.append(LOGO_DIR / Path(logo_fn).name)

    unique: list[Path] = []
    seen: set[str] = set()
    for p in out:
        k = str(p.resolve()) if p.exists() else str(p)
        if k not in seen:
            seen.add(k)
            unique.append(p)
    return unique


def _resolve_logo_path(
    report: dict[str, Any],
    company_profile: dict[str, Any],
    *,
    resolve_logo: LogoPathResolver | None = None,
) -> Path | None:
    if resolve_logo is not None:
        custom = resolve_logo(report, company_profile)
        if custom is not None and custom.is_file():
            return custom
    for p in _candidate_logo_paths(report, company_profile):
        if p.is_file():
            return p
    return None


def _logo_image_for_pdf(path: Path, max_width_cm: float, max_height_cm: float) -> Image | None:
    try:
        iw, ih = ImageReader(str(path)).getSize()
    except Exception:
        return None
    if iw <= 0 or ih <= 0:
        return None

    max_w = max_width_cm * cm
    max_h = max_height_cm * cm
    scale = min(max_w / float(iw), max_h / float(ih), 1.0)
    img = Image(str(path), width=float(iw) * scale, height=float(ih) * scale)
    img.hAlign = "LEFT"
    return img


def _report_signatures_doc(report: dict[str, Any]) -> dict[str, dict[str, Any] | None]:
    raw = report.get("signatures")
    out: dict[str, dict[str, Any] | None] = {"customer": None, "employee": None}
    if not isinstance(raw, dict):
        return out
    for role in SIGNATURE_ROLES:
        entry = raw.get(role)
        if isinstance(entry, dict) and entry.get("filename"):
            out[role] = entry
    return out


def _report_tenant_id(report: dict[str, Any]) -> str | None:
    tid = str(report.get("companyId") or "").strip()
    return tid or None


def _candidate_signature_paths(filename: str, *, tenant_id: str | None = None) -> list[Path]:
    fn = str(filename or "").strip()
    if not fn:
        return []
    out: list[Path] = []
    if tenant_id:
        out.append(TENANTS_UPLOADS_DIR / tenant_id / "signatures" / fn)
    out.append(SIGNATURES_DIR / fn)
    if not tenant_id and TENANTS_UPLOADS_DIR.is_dir():
        for tenant_dir in TENANTS_UPLOADS_DIR.iterdir():
            if tenant_dir.is_dir():
                out.append(tenant_dir / "signatures" / fn)
    unique: list[Path] = []
    seen: set[str] = set()
    for path in out:
        key = str(path)
        if key in seen:
            continue
        seen.add(key)
        unique.append(path)
    return unique


def _safe_signature_path(
    filename: str,
    *,
    resolve_signature: SignaturePathResolver | None = None,
    tenant_id: str | None = None,
) -> Path | None:
    fn = str(filename or "")
    if not fn or "/" in fn or "\\" in fn or fn.strip() != fn:
        return None
    if resolve_signature is not None:
        resolved = resolve_signature(fn)
        if resolved is not None and resolved.is_file():
            return resolved
    for path in _candidate_signature_paths(fn, tenant_id=tenant_id):
        try:
            resolved = path.resolve()
            resolved.relative_to(path.parent.resolve())
        except ValueError:
            continue
        if resolved.is_file():
            return resolved
    return None


def _format_signed_at_de(signed_at: Any, fallback_date: Any = None) -> str:
    s = str(signed_at or "").strip()
    if s:
        m = re.match(r"^(\d{4})-(\d{2})-(\d{2})", s)
        if m:
            return f"{m.group(3)}.{m.group(2)}.{m.group(1)}"
        m2 = re.match(r"^(\d{4})-(\d{1,2})-(\d{1,2})$", s[:10])
        if m2:
            return f"{int(m2.group(3))}.{int(m2.group(2))}.{m2.group(1)}"
    return _format_date_de(str(fallback_date or "—"))


def _signature_image_for_pdf(path: Path, max_width: float, max_height: float) -> Image | None:
    try:
        iw, ih = ImageReader(str(path)).getSize()
    except Exception:
        return None
    if iw <= 0 or ih <= 0:
        return None
    scale = min(max_width / float(iw), max_height / float(ih), 1.0)
    img = Image(str(path), width=float(iw) * scale, height=float(ih) * scale)
    img.hAlign = "CENTER"
    return img


def _pdf_signature_cell(
    entry: dict[str, Any] | None,
    *,
    default_label: str,
    report_date: str,
    col_width: float,
    info_label_style: ParagraphStyle,
    meta_style: ParagraphStyle,
    resolve_signature: SignaturePathResolver | None = None,
    tenant_id: str | None = None,
) -> list[Any]:
    if not isinstance(entry, dict):
        return [Spacer(1, 1)]
    path = _safe_signature_path(
        str(entry.get("filename") or ""),
        resolve_signature=resolve_signature,
        tenant_id=tenant_id,
    )
    if path is None:
        return [Spacer(1, 1)]

    flows: list[Any] = []
    img = _signature_image_for_pdf(path, max(col_width - 16, 40), 2.6 * cm)
    if img:
        flows.append(img)
    else:
        flows.append(Spacer(1, 28))

    label = str(entry.get("signedByLabel") or default_label).strip() or default_label
    date_line = _format_signed_at_de(entry.get("signedAt"), report_date)
    flows.append(Spacer(1, 4))
    flows.append(Paragraph(_xml_para_text(label), info_label_style))
    flows.append(Paragraph(_xml_para_text(f"Datum: {date_line}"), meta_style))
    return flows


def _append_pdf_signatures(
    story: list[Any],
    doc_tpl: SimpleDocTemplate,
    report: dict[str, Any],
    *,
    section_head: ParagraphStyle,
    info_label_style: ParagraphStyle,
    meta_style: ParagraphStyle,
    resolve_signature: SignaturePathResolver | None = None,
) -> None:
    sigs = _report_signatures_doc(report)
    customer = sigs.get("customer")
    employee = sigs.get("employee")
    tenant_id = _report_tenant_id(report)
    has_customer = bool(
        isinstance(customer, dict)
        and _safe_signature_path(
            str(customer.get("filename") or ""),
            resolve_signature=resolve_signature,
            tenant_id=tenant_id,
        )
    )
    has_employee = bool(
        isinstance(employee, dict)
        and _safe_signature_path(
            str(employee.get("filename") or ""),
            resolve_signature=resolve_signature,
            tenant_id=tenant_id,
        )
    )
    if not has_customer and not has_employee:
        return

    report_date = str(report.get("date") or "—")
    col_w = doc_tpl.width * 0.48

    story.append(Spacer(1, 12))
    story.append(Paragraph(_xml_para_text("Unterschriften"), section_head))

    sig_tbl = Table(
        [
            [
                _pdf_signature_cell(
                    customer if has_customer else None,
                    default_label="Kunde",
                    report_date=report_date,
                    col_width=col_w,
                    info_label_style=info_label_style,
                    meta_style=meta_style,
                    resolve_signature=resolve_signature,
                    tenant_id=tenant_id,
                ),
                _pdf_signature_cell(
                    employee if has_employee else None,
                    default_label="Baustellenleitung / Mitarbeiter",
                    report_date=report_date,
                    col_width=col_w,
                    info_label_style=info_label_style,
                    meta_style=meta_style,
                    resolve_signature=resolve_signature,
                    tenant_id=tenant_id,
                ),
            ]
        ],
        colWidths=[col_w, col_w],
        hAlign="LEFT",
    )
    sig_tbl.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("BOX", (0, 0), (0, 0), 0.5, LINE_HEX),
                ("BOX", (1, 0), (1, 0), 0.5, LINE_HEX),
            ]
        )
    )
    story.append(sig_tbl)


def build_pdf_bytes(
    report: dict[str, Any],
    company_profile: dict[str, Any],
    *,
    resolve_logo: LogoPathResolver | None = None,
    resolve_signature: SignaturePathResolver | None = None,
) -> bytes:
    buf = BytesIO()
    doc_tpl = SimpleDocTemplate(
        buf,
        pagesize=A4,
        title="Tagesbericht",
        leftMargin=1.8 * cm,
        rightMargin=1.8 * cm,
        topMargin=1.4 * cm,
        bottomMargin=1.8 * cm,
    )
    styles = getSampleStyleSheet()
    meta_style = ParagraphStyle(
        name="MetaLine",
        parent=styles["Normal"],
        fontSize=9,
        textColor=GREY_META_HEX,
        spaceAfter=1,
        leading=11,
    )
    body_style = ParagraphStyle(
        name="ReportBody",
        parent=styles["Normal"],
        fontSize=9.5,
        leading=13,
        textColor=TEXT_DARK_HEX,
    )

    section_head = ParagraphStyle(
        name="SectionHeadPdfCalm",
        parent=styles["Heading2"],
        fontSize=10.5,
        textColor=SECTION_HEX,
        spaceBefore=10,
        spaceAfter=4,
        leading=13,
        fontName="Helvetica-Bold",
    )

    title_style = ParagraphStyle(
        name="DocTitlePdf",
        parent=styles["Heading1"],
        fontSize=15,
        textColor=TEXT_DARK_HEX,
        spaceBefore=4,
        spaceAfter=6,
        alignment=1,
        fontName="Helvetica-Bold",
    )

    company_style = ParagraphStyle(
        name="CompanyHeadPdf",
        parent=styles["Heading1"],
        fontSize=13,
        textColor=TEXT_DARK_HEX,
        spaceAfter=2,
        fontName="Helvetica-Bold",
    )
    section_text_style = ParagraphStyle(
        name="SectionBodyPdf",
        parent=body_style,
        spaceAfter=4,
    )
    info_label_style = ParagraphStyle(
        name="InfoLabelPdf",
        parent=meta_style,
        fontName="Helvetica-Bold",
        textColor=SECTION_HEX,
    )
    info_value_style = ParagraphStyle(
        name="InfoValuePdf",
        parent=body_style,
        fontSize=9.2,
    )
    bullet_style = ParagraphStyle(
        name="BulletPdf",
        parent=body_style,
        leftIndent=10,
        spaceAfter=3,
    )

    st = _structured(report)
    emails = str(report.get("officeEmail") or company_profile.get("officeEmail") or "")
    phone = str(company_profile.get("phone") or "")
    company_name = str(report.get("companyName") or company_profile.get("companyName") or "Firma")
    proj = str(report.get("projectName") or "—")
    customer = str(report.get("customerName") or "—")
    datum = _format_date_de(str(report.get("date") or "—"))
    emps_raw = report.get("employees")
    mitarbeiter = ", ".join(str(e) for e in emps_raw) if isinstance(emps_raw, list) and emps_raw else "Keine Angabe"
    zeit = f"{report.get('startTime', '?')} – {report.get('endTime', '?')}"

    summary = str(st.get("summary") or "Keine Angabe")
    acts = _list_or_keine(st.get("activities"))
    mats = _list_or_keine(st.get("materials"))
    probs = _list_or_keine(st.get("problems"))
    opens = _list_or_keine(st.get("openItems"))
    ktalk = str(st.get("customerTalk") or "Keine Angabe")

    def footer(canv: Any, __: Any) -> None:
        canv.saveState()
        canv.setFont("Helvetica", 7)
        canv.setFillColor(GREY_META_HEX)
        canv.drawCentredString(A4[0] / 2.0, 1.2 * cm, "Erstellt mit Freiraum Baustellen-Doku")
        canv.restoreState()

    story: list[Any] = []
    logo_path = _resolve_logo_path(report, company_profile, resolve_logo=resolve_logo)
    logo_img = _logo_image_for_pdf(logo_path, max_width_cm=5.0, max_height_cm=2.9) if logo_path else None

    company_lines: list[Any] = [Paragraph(_xml_para_text(company_name), company_style)]
    if emails:
        company_lines.append(Paragraph(_xml_para_text(f"Büro-E-Mail: {emails}"), meta_style))
    if phone:
        company_lines.append(Paragraph(_xml_para_text(f"Telefon: {phone}"), meta_style))

    if logo_img:
        head_tbl = Table(
            [[logo_img, company_lines]],
            colWidths=[doc_tpl.width * 0.30, doc_tpl.width * 0.70],
        )
    else:
        head_tbl = Table(
            [[company_lines]],
            colWidths=[doc_tpl.width],
        )
    head_tbl.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(head_tbl)
    story.append(Spacer(1, 3))
    story.append(Paragraph("TAGESBERICHT", title_style))
    story.append(Spacer(1, 4))
    line_tbl = Table([[""]], colWidths=[doc_tpl.width], rowHeights=[1.2])
    line_tbl.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), LINE_HEX)]))
    story.append(line_tbl)
    story.append(Spacer(1, 8))

    meta_rows = [
        [Paragraph("Baustelle", info_label_style), Paragraph(_xml_para_text(proj), info_value_style)],
        [Paragraph("Kunde", info_label_style), Paragraph(_xml_para_text(customer), info_value_style)],
        [Paragraph("Datum", info_label_style), Paragraph(_xml_para_text(datum), info_value_style)],
        [Paragraph("Mitarbeitende", info_label_style), Paragraph(_xml_para_text(mitarbeiter), info_value_style)],
        [Paragraph("Arbeitszeit", info_label_style), Paragraph(_xml_para_text(zeit), info_value_style)],
    ]
    tbl = Table(meta_rows, colWidths=[doc_tpl.width * 0.30, doc_tpl.width * 0.70])
    tbl.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 0), (-1, -1), SOFT_BG_HEX),
                ("BOX", (0, 0), (-1, -1), 0.5, LINE_HEX),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, LINE_HEX),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(tbl)
    story.append(Spacer(1, 10))

    def sec(title: str) -> Paragraph:
        return Paragraph(_xml_para_text(title), section_head)

    def bullets(items: list[str]) -> None:
        for item in items:
            story.append(Paragraph(f"\u2022 {_xml_para_text(item)}", bullet_style))

    story.append(sec("Zusammenfassung"))
    story.append(Paragraph(_xml_para_text(summary), section_text_style))

    story.append(sec("Tätigkeiten"))
    bullets(acts)

    story.append(sec("Material"))
    bullets(mats)

    story.append(sec("Probleme"))
    bullets(probs)

    story.append(sec("Offene Punkte"))
    bullets(opens)

    story.append(sec("Kundengespräch"))
    story.append(Paragraph(_xml_para_text(ktalk), section_text_style))

    _append_pdf_signatures(
        story,
        doc_tpl,
        report,
        section_head=section_head,
        info_label_style=info_label_style,
        meta_style=meta_style,
        resolve_signature=resolve_signature,
    )

    doc_tpl.build(story, onFirstPage=footer, onLaterPages=footer)
    return buf.getvalue()


META_GREY_DOCX = RGBColor(0x52, 0x52, 0x5B)
TEXT_DARK_DOCX = RGBColor(0x11, 0x18, 0x27)
LINE_GREY_DOCX = RGBColor(0xD4, 0xD4, 0xD8)


def _heading_docx(head: Any) -> None:
    for r in head.runs:
        r.font.color.rgb = TEXT_DARK_DOCX


def _section_list_docx(document: Document, title: str, items: list[str]) -> None:
    h = document.add_heading(title, level=2)
    _heading_docx(h)
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(10)


def build_docx_bytes(
    report: dict[str, Any],
    company_profile: dict[str, Any],
    *,
    resolve_logo: LogoPathResolver | None = None,
) -> bytes:
    st = _structured(report)
    d = Document()

    sec = d.sections[0]
    sec.top_margin = Cm(1.7)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(1.9)
    sec.right_margin = Cm(1.9)

    company_name = str(report.get("companyName") or company_profile.get("companyName") or "Firma")
    email = str(report.get("officeEmail") or company_profile.get("officeEmail") or "")
    phone = str(company_profile.get("phone") or "")
    logo_path = _resolve_logo_path(report, company_profile, resolve_logo=resolve_logo)

    head_tbl = d.add_table(rows=1, cols=2)
    head_tbl.autofit = True
    logo_cell = head_tbl.rows[0].cells[0]
    info_cell = head_tbl.rows[0].cells[1]

    if logo_path:
        p_logo = logo_cell.paragraphs[0]
        p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        try:
            p_logo.add_run().add_picture(str(logo_path), width=Cm(5.0))
        except Exception:
            p_logo.add_run("")
    else:
        logo_cell.text = ""

    p_company = info_cell.paragraphs[0]
    r_company = p_company.add_run(company_name)
    r_company.bold = True
    r_company.font.size = Pt(14)
    r_company.font.color.rgb = TEXT_DARK_DOCX

    if email:
        p_mail = info_cell.add_paragraph(f"Büro-E-Mail: {email}")
        for r in p_mail.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = META_GREY_DOCX
    if phone:
        p_ph = info_cell.add_paragraph(f"Telefon: {phone}")
        for r in p_ph.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = META_GREY_DOCX

    d.add_paragraph()
    p_title = d.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r_title = p_title.add_run("TAGESBERICHT")
    r_title.bold = True
    r_title.font.size = Pt(15)
    r_title.font.color.rgb = TEXT_DARK_DOCX

    p_div = d.add_paragraph("_______________________________________________")
    p_div.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for r in p_div.runs:
        r.font.color.rgb = LINE_GREY_DOCX
        r.font.size = Pt(8)

    proj = str(report.get("projectName") or "—")
    customer = str(report.get("customerName") or "—")
    datum = _format_date_de(str(report.get("date") or "—"))
    emps_raw = report.get("employees")
    mitarbeiter = ", ".join(str(e) for e in emps_raw) if isinstance(emps_raw, list) and emps_raw else "Keine Angabe"
    zeit = f"{report.get('startTime', '?')} – {report.get('endTime', '?')}"

    info_tbl = d.add_table(rows=0, cols=2)
    info_tbl.style = "Table Grid"
    for label, value in [
        ("Baustelle", proj),
        ("Kunde", customer),
        ("Datum", datum),
        ("Mitarbeitende", mitarbeiter),
        ("Arbeitszeit", zeit),
    ]:
        row = info_tbl.add_row().cells
        p_l = row[0].paragraphs[0]
        r_l = p_l.add_run(label)
        r_l.bold = True
        r_l.font.color.rgb = META_GREY_DOCX
        r_l.font.size = Pt(10)

        p_v = row[1].paragraphs[0]
        r_v = p_v.add_run(value)
        r_v.font.size = Pt(10)
        r_v.font.color.rgb = TEXT_DARK_DOCX

    d.add_paragraph()

    summary = str(st.get("summary") or "Keine Angabe")
    acts = _list_or_keine(st.get("activities"))
    mats = _list_or_keine(st.get("materials"))
    probs = _list_or_keine(st.get("problems"))
    opens = _list_or_keine(st.get("openItems"))
    ktalk = str(st.get("customerTalk") or "Keine Angabe")

    h_sum = d.add_heading("Zusammenfassung", level=2)
    _heading_docx(h_sum)
    p_sum = d.add_paragraph(summary)
    for r in p_sum.runs:
        r.font.size = Pt(10)

    _section_list_docx(d, "Tätigkeiten", acts)
    _section_list_docx(d, "Material", mats)
    _section_list_docx(d, "Probleme", probs)
    _section_list_docx(d, "Offene Punkte", opens)

    h_talk = d.add_heading("Kundengespräch", level=2)
    _heading_docx(h_talk)
    p_talk = d.add_paragraph(ktalk)
    for r in p_talk.runs:
        r.font.size = Pt(10)

    d.add_paragraph()
    foot = d.add_paragraph("Erstellt mit Freiraum Baustellen-Doku")
    foot.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for r in foot.runs:
        r.font.size = Pt(8)
        r.italic = True
        r.font.color.rgb = META_GREY_DOCX

    bio = BytesIO()
    d.save(bio)
    return bio.getvalue()


# ============================================================================
# Gesamtbericht / Sammelbericht (Folgebericht) — eigenstaendige Builder.
# Bestehende Tagesbericht-Builder bleiben unveraendert.
# ============================================================================

PhotoPathResolver = Callable[[str], Path | None]

_MAX_COLLECTIVE_PHOTOS = 24


def collective_export_base_name(payload: dict[str, Any]) -> str:
    site = sanitize_export_slug(str(payload.get("projectName") or "Baustelle"))
    return f"gesamtbericht_{site}"


def build_collective_attachment_names(payload: dict[str, Any], ext: str) -> tuple[str, str]:
    ascii_nm = (
        "gesamtbericht_"
        f"{sanitize_export_slug_ascii(str(payload.get('projectName') or 'Baustelle')).lower()}.{ext}"
    )
    desc = f"{collective_export_base_name(payload)}.{ext}"
    return ascii_nm, desc


def _zeitraum_label(payload: dict[str, Any]) -> str:
    df = _format_date_de(str(payload.get("dateFrom") or ""))
    dt = _format_date_de(str(payload.get("dateTo") or ""))
    if df and dt and df != dt:
        return f"{df} – {dt}"
    return df or dt or "—"


def _fmt_hours(value: Any) -> str:
    try:
        return f"{float(value):.2f}".replace(".", ",")
    except Exception:
        return "0,00"


def build_collective_pdf_bytes(
    payload: dict[str, Any],
    company_profile: dict[str, Any],
    *,
    resolve_logo: LogoPathResolver | None = None,
    resolve_photo: PhotoPathResolver | None = None,
    resolve_signature: SignaturePathResolver | None = None,
) -> bytes:
    buf = BytesIO()
    doc_tpl = SimpleDocTemplate(
        buf,
        pagesize=A4,
        title="Gesamtbericht",
        leftMargin=1.8 * cm,
        rightMargin=1.8 * cm,
        topMargin=1.4 * cm,
        bottomMargin=1.8 * cm,
    )
    styles = getSampleStyleSheet()
    meta_style = ParagraphStyle("MetaLineC", parent=styles["Normal"], fontSize=9, textColor=GREY_META_HEX, spaceAfter=1, leading=11)
    body_style = ParagraphStyle("ReportBodyC", parent=styles["Normal"], fontSize=9.5, leading=13, textColor=TEXT_DARK_HEX)
    section_head = ParagraphStyle("SectionHeadC", parent=styles["Heading2"], fontSize=10.5, textColor=SECTION_HEX, spaceBefore=10, spaceAfter=4, leading=13, fontName="Helvetica-Bold")
    day_head = ParagraphStyle("DayHeadC", parent=styles["Heading3"], fontSize=9.8, textColor=SECTION_HEX, spaceBefore=8, spaceAfter=2, leading=12, fontName="Helvetica-Bold")
    title_style = ParagraphStyle("DocTitleC", parent=styles["Heading1"], fontSize=15, textColor=TEXT_DARK_HEX, spaceBefore=4, spaceAfter=6, alignment=1, fontName="Helvetica-Bold")
    company_style = ParagraphStyle("CompanyHeadC", parent=styles["Heading1"], fontSize=13, textColor=TEXT_DARK_HEX, spaceAfter=2, fontName="Helvetica-Bold")
    section_text_style = ParagraphStyle("SectionBodyC", parent=body_style, spaceAfter=4)
    info_label_style = ParagraphStyle("InfoLabelC", parent=meta_style, fontName="Helvetica-Bold", textColor=SECTION_HEX)
    info_value_style = ParagraphStyle("InfoValueC", parent=body_style, fontSize=9.2)
    bullet_style = ParagraphStyle("BulletC", parent=body_style, leftIndent=10, spaceAfter=3)
    note_style = ParagraphStyle("NoteC", parent=body_style, leftIndent=10, spaceAfter=3, textColor=GREY_META_HEX)

    company_name = str(payload.get("companyName") or company_profile.get("companyName") or "Firma")
    emails = str(payload.get("officeEmail") or company_profile.get("officeEmail") or "")
    phone = str(company_profile.get("phone") or "")

    def footer(canv: Any, __: Any) -> None:
        canv.saveState()
        canv.setFont("Helvetica", 7)
        canv.setFillColor(GREY_META_HEX)
        canv.drawCentredString(A4[0] / 2.0, 1.2 * cm, "Erstellt mit Freiraum Baustellen-Doku")
        canv.restoreState()

    story: list[Any] = []
    logo_path = _resolve_logo_path(payload, company_profile, resolve_logo=resolve_logo)
    logo_img = _logo_image_for_pdf(logo_path, max_width_cm=5.0, max_height_cm=2.9) if logo_path else None

    company_lines: list[Any] = [Paragraph(_xml_para_text(company_name), company_style)]
    if emails:
        company_lines.append(Paragraph(_xml_para_text(f"Büro-E-Mail: {emails}"), meta_style))
    if phone:
        company_lines.append(Paragraph(_xml_para_text(f"Telefon: {phone}"), meta_style))

    if logo_img:
        head_tbl = Table([[logo_img, company_lines]], colWidths=[doc_tpl.width * 0.30, doc_tpl.width * 0.70])
    else:
        head_tbl = Table([[company_lines]], colWidths=[doc_tpl.width])
    head_tbl.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0), ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
    story.append(head_tbl)
    story.append(Spacer(1, 3))
    story.append(Paragraph("GESAMTBERICHT", title_style))
    story.append(Spacer(1, 4))
    line_tbl = Table([[""]], colWidths=[doc_tpl.width], rowHeights=[1.2])
    line_tbl.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), LINE_HEX)]))
    story.append(line_tbl)
    story.append(Spacer(1, 8))

    totals = payload.get("totals", {}) if isinstance(payload.get("totals"), dict) else {}
    meta_rows = [
        [Paragraph("Baustelle", info_label_style), Paragraph(_xml_para_text(str(payload.get("projectName") or "—")), info_value_style)],
        [Paragraph("Kunde", info_label_style), Paragraph(_xml_para_text(str(payload.get("customerName") or "—")), info_value_style)],
        [Paragraph("Zeitraum", info_label_style), Paragraph(_xml_para_text(_zeitraum_label(payload)), info_value_style)],
        [Paragraph("Arbeitstage", info_label_style), Paragraph(_xml_para_text(str(totals.get("reportCount") or 0)), info_value_style)],
        [Paragraph("Gesamtstunden", info_label_style), Paragraph(_xml_para_text(_fmt_hours(totals.get("totalHours"))), info_value_style)],
    ]
    tbl = Table(meta_rows, colWidths=[doc_tpl.width * 0.30, doc_tpl.width * 0.70])
    tbl.setStyle(TableStyle([("FONTNAME", (0, 0), (-1, -1), "Helvetica"), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("BACKGROUND", (0, 0), (-1, -1), SOFT_BG_HEX), ("BOX", (0, 0), (-1, -1), 0.5, LINE_HEX), ("INNERGRID", (0, 0), (-1, -1), 0.25, LINE_HEX), ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6), ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
    story.append(tbl)
    story.append(Spacer(1, 10))

    def sec(title: str) -> None:
        story.append(Paragraph(_xml_para_text(title), section_head))

    def bullets(items: list[str], style: ParagraphStyle = bullet_style) -> None:
        for item in items:
            story.append(Paragraph(f"\u2022 {_xml_para_text(str(item))}", style))

    sec("Zusammenfassung")
    story.append(Paragraph(_xml_para_text(str(payload.get("summary") or "Keine Angabe")), section_text_style))

    hours_by_emp = totals.get("hoursByEmployee") or []
    if hours_by_emp:
        sec("Stunden je Mitarbeiter")
        for row in hours_by_emp:
            if isinstance(row, dict):
                story.append(Paragraph(f"\u2022 {_xml_para_text(str(row.get('name') or '—'))}: {_fmt_hours(row.get('hours'))} h", bullet_style))

    mats = totals.get("materials") or []
    if mats:
        sec("Material (gesamt)")
        bullets([str(m) for m in mats])

    opens = totals.get("openItems") or []
    if opens:
        sec("Offene Punkte (gesamt)")
        bullets([str(o) for o in opens])

    probs = totals.get("problems") or []
    if probs:
        sec("Probleme (gesamt)")
        bullets([str(p) for p in probs])

    days = payload.get("days") or []
    if days:
        sec("Tagesverlauf")
        for day in days:
            if not isinstance(day, dict):
                continue
            emps = day.get("employees") or []
            emps_label = ", ".join(str(e) for e in emps) if emps else "—"
            head = f"{_format_date_de(str(day.get('date') or '—'))}  ·  {emps_label}  ·  {_fmt_hours(day.get('hours'))} h"
            story.append(Paragraph(_xml_para_text(head), day_head))
            bullets([str(a) for a in (day.get("activities") or [])])
            note = str(day.get("notes") or "").strip()
            if note:
                story.append(Paragraph(f"Besonderheiten: {_xml_para_text(note)}", note_style))

    signatures = payload.get("signatures") or []
    if signatures:
        rendered: list[Any] = []
        role_label = {"customer": "Kunde / Auftraggeber", "employee": "Mitarbeiter"}
        for sg in signatures:
            if not isinstance(sg, dict):
                continue
            fn = str(sg.get("filename") or "")
            path = _safe_signature_path(fn, resolve_signature=resolve_signature)
            if path is None:
                continue
            img = _signature_image_for_pdf(path, 6.5 * cm, 2.4 * cm)
            if img is None:
                continue
            label = str(sg.get("signedByLabel") or role_label.get(str(sg.get("role")), "Unterschrift")).strip()
            date_line = _format_signed_at_de(sg.get("signedAt"), sg.get("date"))
            rendered.append(Spacer(1, 4))
            rendered.append(img)
            rendered.append(Paragraph(_xml_para_text(label), info_label_style))
            rendered.append(Paragraph(_xml_para_text(f"Datum: {date_line}"), meta_style))
        if rendered:
            sec("Unterschriften")
            for flow in rendered:
                story.append(flow)

    photos = payload.get("photos") or []
    if photos and resolve_photo is not None:
        sec("Fotos")
        shown = 0
        for ph in photos:
            if shown >= _MAX_COLLECTIVE_PHOTOS:
                break
            if not isinstance(ph, dict):
                continue
            fn = ph.get("filename")
            if not isinstance(fn, str) or not fn:
                continue
            path = resolve_photo(fn)
            if path is None:
                continue
            img = _logo_image_for_pdf(path, max_width_cm=8.0, max_height_cm=6.0)
            if img is None:
                continue
            story.append(Spacer(1, 4))
            story.append(Paragraph(_xml_para_text(f"{_format_date_de(str(ph.get('date') or ''))}"), meta_style))
            story.append(img)
            shown += 1

    doc_tpl.build(story, onFirstPage=footer, onLaterPages=footer)
    return buf.getvalue()


def build_collective_docx_bytes(
    payload: dict[str, Any],
    company_profile: dict[str, Any],
    *,
    resolve_logo: LogoPathResolver | None = None,
    resolve_photo: PhotoPathResolver | None = None,
    resolve_signature: SignaturePathResolver | None = None,
) -> bytes:
    d = Document()
    sec = d.sections[0]
    sec.top_margin = Cm(1.7)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(1.9)
    sec.right_margin = Cm(1.9)

    company_name = str(payload.get("companyName") or company_profile.get("companyName") or "Firma")
    email = str(payload.get("officeEmail") or company_profile.get("officeEmail") or "")
    phone = str(company_profile.get("phone") or "")
    logo_path = _resolve_logo_path(payload, company_profile, resolve_logo=resolve_logo)

    head_tbl = d.add_table(rows=1, cols=2)
    head_tbl.autofit = True
    logo_cell = head_tbl.rows[0].cells[0]
    info_cell = head_tbl.rows[0].cells[1]
    if logo_path:
        try:
            logo_cell.paragraphs[0].add_run().add_picture(str(logo_path), width=Cm(5.0))
        except Exception:
            logo_cell.text = ""
    else:
        logo_cell.text = ""
    r_company = info_cell.paragraphs[0].add_run(company_name)
    r_company.bold = True
    r_company.font.size = Pt(14)
    r_company.font.color.rgb = TEXT_DARK_DOCX
    if email:
        for r in info_cell.add_paragraph(f"Büro-E-Mail: {email}").runs:
            r.font.size = Pt(9)
            r.font.color.rgb = META_GREY_DOCX
    if phone:
        for r in info_cell.add_paragraph(f"Telefon: {phone}").runs:
            r.font.size = Pt(9)
            r.font.color.rgb = META_GREY_DOCX

    d.add_paragraph()
    p_title = d.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r_title = p_title.add_run("GESAMTBERICHT")
    r_title.bold = True
    r_title.font.size = Pt(15)
    r_title.font.color.rgb = TEXT_DARK_DOCX

    totals = payload.get("totals", {}) if isinstance(payload.get("totals"), dict) else {}
    info_tbl = d.add_table(rows=0, cols=2)
    info_tbl.style = "Table Grid"
    for label, value in [
        ("Baustelle", str(payload.get("projectName") or "—")),
        ("Kunde", str(payload.get("customerName") or "—")),
        ("Zeitraum", _zeitraum_label(payload)),
        ("Arbeitstage", str(totals.get("reportCount") or 0)),
        ("Gesamtstunden", _fmt_hours(totals.get("totalHours"))),
    ]:
        row = info_tbl.add_row().cells
        r_l = row[0].paragraphs[0].add_run(label)
        r_l.bold = True
        r_l.font.color.rgb = META_GREY_DOCX
        r_l.font.size = Pt(10)
        r_v = row[1].paragraphs[0].add_run(value)
        r_v.font.size = Pt(10)
        r_v.font.color.rgb = TEXT_DARK_DOCX

    d.add_paragraph()
    h_sum = d.add_heading("Zusammenfassung", level=2)
    _heading_docx(h_sum)
    for r in d.add_paragraph(str(payload.get("summary") or "Keine Angabe")).runs:
        r.font.size = Pt(10)

    hours_by_emp = totals.get("hoursByEmployee") or []
    if hours_by_emp:
        _section_list_docx(d, "Stunden je Mitarbeiter", [f"{row.get('name')}: {_fmt_hours(row.get('hours'))} h" for row in hours_by_emp if isinstance(row, dict)])
    if totals.get("materials"):
        _section_list_docx(d, "Material (gesamt)", [str(m) for m in totals.get("materials")])
    if totals.get("openItems"):
        _section_list_docx(d, "Offene Punkte (gesamt)", [str(o) for o in totals.get("openItems")])
    if totals.get("problems"):
        _section_list_docx(d, "Probleme (gesamt)", [str(p) for p in totals.get("problems")])

    days = payload.get("days") or []
    if days:
        h_days = d.add_heading("Tagesverlauf", level=2)
        _heading_docx(h_days)
        for day in days:
            if not isinstance(day, dict):
                continue
            emps = day.get("employees") or []
            emps_label = ", ".join(str(e) for e in emps) if emps else "—"
            h_day = d.add_heading(f"{_format_date_de(str(day.get('date') or '—'))} · {emps_label} · {_fmt_hours(day.get('hours'))} h", level=3)
            _heading_docx(h_day)
            for a in (day.get("activities") or []):
                p = d.add_paragraph(style="List Bullet")
                p.add_run(str(a)).font.size = Pt(10)
            note = str(day.get("notes") or "").strip()
            if note:
                for r in d.add_paragraph(f"Besonderheiten: {note}").runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = META_GREY_DOCX

    signatures = payload.get("signatures") or []
    if signatures:
        role_label = {"customer": "Kunde / Auftraggeber", "employee": "Mitarbeiter"}
        sig_rendered = False
        for sg in signatures:
            if not isinstance(sg, dict):
                continue
            path = _safe_signature_path(str(sg.get("filename") or ""), resolve_signature=resolve_signature)
            if path is None:
                continue
            if not sig_rendered:
                _heading_docx(d.add_heading("Unterschriften", level=2))
                sig_rendered = True
            try:
                d.add_paragraph().add_run().add_picture(str(path), width=Cm(6.5))
            except Exception:
                continue
            label = str(sg.get("signedByLabel") or role_label.get(str(sg.get("role")), "Unterschrift")).strip()
            date_line = _format_signed_at_de(sg.get("signedAt"), sg.get("date"))
            for r in d.add_paragraph(f"{label} — Datum: {date_line}").runs:
                r.font.size = Pt(9)
                r.font.color.rgb = META_GREY_DOCX

    photos = payload.get("photos") or []
    if photos and resolve_photo is not None:
        h_ph = d.add_heading("Fotos", level=2)
        _heading_docx(h_ph)
        shown = 0
        for ph in photos:
            if shown >= _MAX_COLLECTIVE_PHOTOS:
                break
            if not isinstance(ph, dict):
                continue
            fn = ph.get("filename")
            if not isinstance(fn, str) or not fn:
                continue
            path = resolve_photo(fn)
            if path is None:
                continue
            try:
                d.add_paragraph(f"{_format_date_de(str(ph.get('date') or ''))}").runs[0].font.size = Pt(8)
                d.add_paragraph().add_run().add_picture(str(path), width=Cm(8.0))
                shown += 1
            except Exception:
                continue

    d.add_paragraph()
    foot = d.add_paragraph("Erstellt mit Freiraum Baustellen-Doku")
    foot.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for r in foot.runs:
        r.font.size = Pt(8)
        r.italic = True
        r.font.color.rgb = META_GREY_DOCX

    bio = BytesIO()
    d.save(bio)
    return bio.getvalue()
